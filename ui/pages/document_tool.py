import json
import sys
import subprocess
import tempfile
import importlib
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from PySide6.QtCore import QDate, QRegularExpression, QStandardPaths, Qt, QUrl
from PySide6.QtGui import QDesktopServices, QRegularExpressionValidator
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTableWidget, QTableWidgetItem, 
    QPushButton, QDialog, QLabel, QFormLayout, QLineEdit, QMessageBox, QHeaderView, QScrollArea,
    QApplication, QComboBox, QCheckBox, QDateEdit, QFileDialog, QProgressDialog, QStackedWidget
)
from ui.pages.logging_util import log


FIELD_TYPES = {
    "text": "텍스트",
    "branch": "분기",
    "date": "날짜",
    "amount": "숫자",
    "branch_value": "지점의 값",
    "table_list": "테이블 목록",
    "folder": "폴더의 이미지",
}


def branch_control_parts(field_key):
    """Return (group name, option name) for 분기_<field>_<option> controls."""
    if not isinstance(field_key, str):
        return None
    parts = field_key.split("_", 2)
    if len(parts) != 3 or parts[0] != "분기" or not parts[1].strip() or not parts[2].strip():
        return None
    return parts[1].strip(), parts[2].strip()

IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}

TABLE_LIST_TYPE_PREFIX = "table_list:"


def table_list_type(table_id):
    return f"{TABLE_LIST_TYPE_PREFIX}{table_id}"


def table_list_id(field_type):
    if isinstance(field_type, str) and field_type.startswith(TABLE_LIST_TYPE_PREFIX):
        return field_type[len(TABLE_LIST_TYPE_PREFIX):]
    return ""


def table_list_options(table):
    options = []
    for row in table.get("cells", []):
        if not isinstance(row, list):
            continue
        options.extend(str(value).strip() for value in row if str(value).strip())
    return options

BRANCH_SOURCE_TYPES = {
    "branch_select": "지점 목록",
    "branch_select_2": "지점 목록 2",
}

DATE_FORMATS = {
    "yyyy-MM-dd": "yyyy-mm-dd",
    "yyyy년 MM월 dd일": "yyyy년 mm월 dd일",
    "yyyy.MM.dd": "yyyy.mm.dd",
}

STANDARD_BRANCH_VALUE_FIELDS = {
    "organization_code": "기관기호",
    "organization_name": "기관명",
    "branch_name": "지점명",
    "corporation_name": "법인명",
    "owner_name": "대표자명",
    "address": "주소",
}


def read_json_file(path, default):
    if not path.exists():
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def write_json_file(path, data):
    path.parent.mkdir(exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def import_openpyxl():
    try:
        return importlib.import_module("openpyxl")
    except ImportError:
        return None


class DocumentTool(QWidget):
    def __init__(self):
        super().__init__()
        
        self.document_folder = Path(__file__).resolve().parents[2] / getattr(self, "document_folder_name", "document")
        self.document_pattern = getattr(self, "document_pattern", "*.docx")
        self.tool_title = getattr(self, "tool_title", "문서 작성 도구(워드)")
        self.document_folder.mkdir(exist_ok=True)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)
        
        title_layout = QHBoxLayout()
        title = QLabel(self.tool_title)
        title.setStyleSheet("font-weight: bold; font-size: 14px;")
        open_folder_button = QPushButton("폴더 열기")
        open_folder_button.clicked.connect(self.open_document_folder)
        title_layout.addWidget(title)
        title_layout.addStretch()
        title_layout.addWidget(open_folder_button)
        layout.addLayout(title_layout)
        
        # 테이블
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["파일명", "설정", "생성", "엑셀다운", "엑셀업로드"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        for column in range(1, 5):
            self.table.horizontalHeader().setSectionResizeMode(column, QHeaderView.ResizeToContents)
        layout.addWidget(self.table)
        
        self.setLayout(layout)
        self.load_documents()
    
    def load_documents(self):
        """문서 폴더에서 이 도구가 지원하는 파일을 로드한다."""
        self.table.setRowCount(0)
        
        if not self.document_folder.exists():
            return
        
        docx_files = list(self.document_folder.glob(self.document_pattern))
        self.table.setRowCount(len(docx_files))
        
        for row, file_path in enumerate(docx_files):
            # 파일명
            filename_item = QTableWidgetItem(file_path.name)
            filename_item.setFlags(filename_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.table.setItem(row, 0, filename_item)
            
            settings_btn = QPushButton("설정")
            settings_btn.clicked.connect(lambda _, f=file_path: self.open_document_settings(f))
            self.table.setCellWidget(row, 1, settings_btn)

            create_btn = QPushButton("생성")
            create_btn.clicked.connect(lambda _, f=file_path: self.open_document_generation(f))
            self.table.setCellWidget(row, 2, create_btn)

            download_btn = QPushButton("엑셀다운")
            download_btn.clicked.connect(lambda _, f=file_path: self.download_excel_template(f))
            self.table.setCellWidget(row, 3, download_btn)

            upload_btn = QPushButton("엑셀업로드")
            upload_btn.clicked.connect(lambda _, f=file_path: self.upload_excel_values(f))
            self.table.setCellWidget(row, 4, upload_btn)

    def open_document_folder(self):
        if not QDesktopServices.openUrl(QUrl.fromLocalFile(str(self.document_folder))):
            QMessageBox.warning(self, "오류", f"문서 폴더를 열 수 없습니다:\n{self.document_folder}")
    
    def open_document_settings(self, file_path):
        """문서 설정 팝업 열기"""
        log("DocumentTool", f"문서 설정 열기 시도: {file_path.name}", level="INFO")
        try:
            dialog = self._create_dialog(file_path, mode="settings")
            dialog.exec()
            self.load_documents()
            log("DocumentTool", f"문서 설정 열기 완료: {file_path.name}", level="INFO")
        except Exception as e:
            log("DocumentTool", f"문서 설정 열기 실패: {e}", level="ERROR")
            QMessageBox.critical(self, "오류", f"문서를 열 수 없습니다: {e}")

    def open_document_generation(self, file_path):
        """문서 생성 팝업 열기"""
        log("DocumentTool", f"문서 생성 열기 시도: {file_path.name}", level="INFO")
        try:
            dialog = self._create_dialog(file_path, mode="generate")
            dialog.exec()
            log("DocumentTool", f"문서 생성 창 닫힘: {file_path.name}", level="INFO")
        except Exception as e:
            log("DocumentTool", f"문서 생성 열기 실패: {e}", level="ERROR")
            QMessageBox.critical(self, "오류", f"문서를 열 수 없습니다: {e}")

    def _settings_path(self):
        filename = getattr(self, "settings_filename", "document_field_settings.json")
        return Path(__file__).resolve().parents[2] / "data" / filename

    def _create_dialog(self, file_path, mode):
        doc = Document(file_path)
        return ContentControlDialog(self, doc, file_path, mode=mode)

    def _load_document_settings(self, file_path):
        settings = read_json_file(self._settings_path(), {"version": 1, "fields": {}, "documents": {}})
        return ContentControlDialog.settings_for_file(settings, file_path)

    def _configured_field_names(self, file_path):
        settings = self._load_document_settings(file_path)
        return [
            field_name
            for field_name, field_setting in settings.get("fields", {}).items()
            if field_setting.get("type") != "folder"
            and not (
                field_setting.get("type") == "date"
                and field_setting.get("date_default_type") == "field_calculation"
            )
        ]

    def _excel_headers(self, file_path):
        settings = self._load_document_settings(file_path)
        tables = read_json_file(self.document_folder.parent / "data" / "table_lists.json", {}).get("tables", [])
        table_names = {table.get("id", ""): table.get("name", "테이블 목록") for table in tables}
        def field_type_name(field_setting):
            field_type = field_setting.get("type", "text")
            target_id = field_setting.get("table_list_id") or table_list_id(field_type)
            if field_type == "table_list" or target_id:
                return table_names.get(target_id, "테이블 목록")
            return FIELD_TYPES.get(field_type, "텍스트")
        return [
            f"{field_name}({field_type_name(field_setting)})"
            for field_name, field_setting in settings.get("fields", {}).items()
            if field_setting.get("type") != "folder"
            and not (
                field_setting.get("type") == "date"
                and field_setting.get("date_default_type") == "field_calculation"
            )
        ]

    def _document_control_names(self, file_path):
        dialog = self._create_dialog(file_path, mode="batch")
        return list(dialog.controls.keys())

    def _excel_sample_values(self, file_path):
        field_names = self._configured_field_names(file_path)
        dialog = self._create_dialog(file_path, mode="batch")
        sample_values = []
        for field_name in field_names:
            if field_name not in dialog.generate_widgets:
                sample_values.append("")
                continue
            value = dialog._field_value(field_name)
            if isinstance(value, dict):
                value = value.get("path", "")
            sample_values.append(value)
        return sample_values

    def download_excel_template(self, file_path):
        openpyxl = import_openpyxl()
        if openpyxl is None:
            QMessageBox.warning(self, "오류", "openpyxl 패키지가 설치되어 있지 않아 엑셀 파일을 만들 수 없습니다.")
            return

        field_names = self._configured_field_names(file_path)
        if not field_names:
            QMessageBox.warning(self, "설정 필요", "먼저 문서의 콘텐츠 컨트롤 필드를 설정해 주세요.")
            return
        missing_fields = [field for field in field_names if field not in self._document_control_names(file_path)]
        if missing_fields:
            QMessageBox.warning(
                self,
                "설정 확인 필요",
                "설정된 필드가 현재 Word 문서에 없습니다:\n" + "\n".join(missing_fields),
            )
            return

        excel_headers = self._excel_headers(file_path)
        default_name = f"{file_path.stem}_업로드양식.xlsx"
        downloads_folder = QStandardPaths.writableLocation(QStandardPaths.DownloadLocation)
        default_path = str(Path(downloads_folder) / default_name) if downloads_folder else default_name
        save_path, _ = QFileDialog.getSaveFileName(self, "엑셀 양식 저장", default_path, "Excel Files (*.xlsx)")
        if not save_path:
            return
        if not save_path.lower().endswith(".xlsx"):
            save_path += ".xlsx"

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "문서생성"
        ws.append(excel_headers)
        ws.append(self._excel_sample_values(file_path))
        for column_index, header in enumerate(excel_headers, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(column_index)].width = max(14, len(header) + 4)
        wb.save(save_path)
        QMessageBox.information(self, "완료", f"엑셀 양식을 저장했습니다:\n{save_path}")

    def upload_excel_values(self, file_path):
        openpyxl = import_openpyxl()
        if openpyxl is None:
            QMessageBox.warning(self, "오류", "openpyxl 패키지가 설치되어 있지 않아 엑셀 파일을 읽을 수 없습니다.")
            return

        expected_fields = self._configured_field_names(file_path)
        expected_headers = self._excel_headers(file_path)
        if not expected_fields:
            QMessageBox.warning(self, "설정 필요", "먼저 문서의 콘텐츠 컨트롤 필드를 설정해 주세요.")
            return
        missing_fields = [field for field in expected_fields if field not in self._document_control_names(file_path)]
        if missing_fields:
            QMessageBox.warning(
                self,
                "설정 확인 필요",
                "설정된 필드가 현재 Word 문서에 없습니다:\n" + "\n".join(missing_fields),
            )
            return

        upload_path, _ = QFileDialog.getOpenFileName(self, "엑셀 업로드", "", "Excel Files (*.xlsx)")
        if not upload_path:
            return

        try:
            wb = openpyxl.load_workbook(upload_path, data_only=True)
            ws = wb.active
            headers = [str(cell.value or "").strip() for cell in ws[1]]
            if headers != expected_headers:
                QMessageBox.warning(
                    self,
                    "양식 불일치",
                    "엑셀 컬럼이 현재 문서 설정과 다릅니다.\n엑셀다운으로 새 양식을 받은 뒤 입력해 주세요.",
                )
                return

            rows = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if not any(value not in (None, "") for value in row):
                    continue
                rows.append({field: "" if value is None else str(value) for field, value in zip(expected_fields, row)})

            if not rows:
                QMessageBox.warning(self, "데이터 없음", "생성할 데이터 행이 없습니다.")
                return

            dialog = self._create_dialog(file_path, mode="batch")
            outputs = dialog.create_documents_from_rows(rows)
            QMessageBox.information(self, "완료", f"{len(outputs)}개 문서를 생성했습니다.")
        except Exception as e:
            log("DocumentTool", f"엑셀 업로드 생성 실패: {e}", level="ERROR")
            QMessageBox.critical(self, "오류", f"엑셀 업로드 처리 중 오류가 발생했습니다:\n{e}")
    
    def showEvent(self, event):
        """페이지가 보여질 때마다 새로고침"""
        super().showEvent(event)
        self.load_documents()



class ContentControlDialog(QDialog):
    @staticmethod
    def settings_for_file(settings, file_path):
        file_key = file_path.name
        documents = settings.get("documents", {})
        if isinstance(documents, dict) and file_key in documents:
            return documents[file_key]
        return {
            "version": 1,
            "branches": settings.get("branches", {}),
            "fields": settings.get("fields", {}),
        }

    def __init__(self, parent, doc, file_path, mode="settings"):
        super().__init__(parent)
        self.mode = mode
        title_action = "설정" if mode == "settings" else "생성"
        self.setWindowTitle(f"{file_path.stem} {title_action}")
        self.setMinimumWidth(760)
        self.setMinimumHeight(600)
        
        self.doc = doc
        self.file_path = file_path
        self.root = Path(__file__).resolve().parents[2]
        self.settings_path = self.root / "data" / getattr(
            self,
            "settings_filename",
            "document_field_settings.json",
        )
        self.branches_path = self.root / "data" / "branches.json"
        self.branch_fields_path = self.root / "data" / "branch_fields.json"
        self.table_lists_path = self.root / "data" / "table_lists.json"
        self.all_settings = read_json_file(self.settings_path, {"version": 1, "fields": {}, "documents": {}})
        self.settings = self.settings_for_file(self.all_settings, file_path)
        self.branches = [
            branch
            for branch in read_json_file(self.branches_path, [])
            if branch.get("active", True)
        ]
        self.branch_fields = read_json_file(self.branch_fields_path, [])
        self.table_lists = read_json_file(self.table_lists_path, {"version": 1, "tables": []}).get("tables", [])
        self._migrate_branch_source_settings()
        self.controls = {}
        self.field_widgets = {}
        self.generate_widgets = {}
        self.title_edit = None
        self.branch_combos = {}
        
        layout = QVBoxLayout()

        branch_title = QLabel("<b>지점 선택</b>")
        layout.addWidget(branch_title)
        branch_layout = QFormLayout()
        for source_type, label in BRANCH_SOURCE_TYPES.items():
            combo = QComboBox()
            combo.setMaximumWidth(240)
            for branch in self.branches:
                combo.addItem(self._branch_name(branch), branch)
            preferred_name = self.settings.get("branches", {}).get(source_type, "")
            self._select_branch_combo(combo, preferred_name)
            combo.currentIndexChanged.connect(self._on_global_branch_changed)
            self.branch_combos[source_type] = combo
            branch_layout.addRow(label, combo)
        layout.addLayout(branch_layout)

        form_layout = QFormLayout()
        form_layout.setRowWrapPolicy(QFormLayout.WrapLongRows)
        
        collected_controls = self._collect_controls(doc)
        # The naming convention is the source of truth. This also makes newly
        # added HWP fields appear as a branch without requiring a first save.
        for field_key in collected_controls:
            if branch_control_parts(field_key):
                self.settings.setdefault("fields", {}).setdefault(field_key, {})["type"] = "branch"
        self.controls.update(collected_controls)
        if self.mode == "settings":
            rendered_groups = set()
            for field_key, control in collected_controls.items():
                parts = branch_control_parts(field_key)
                if parts:
                    if parts[0] not in rendered_groups:
                        members = [key for key in collected_controls if branch_control_parts(key) and branch_control_parts(key)[0] == parts[0]]
                        self._add_branch_settings_group(form_layout, parts[0], members)
                        rendered_groups.add(parts[0])
                    continue
                self._add_field_row(form_layout, field_key, control)
        if self.mode == "settings":
            self._refresh_type_combo_options()
            self._refresh_branch_value_source_options()
        else:
            if self.mode == "generate":
                self.title_edit = QLineEdit(self.file_path.stem)
                self.title_edit.setPlaceholderText("저장할 문서 제목")
                form_layout.addRow("저장 제목", self.title_edit)
                form_layout.addRow("", QLabel(""))
            self._add_generate_rows(form_layout)
        self._update_calculated_date_fields()

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        container = QWidget()
        container.setLayout(form_layout)
        scroll.setWidget(container)

        layout.addWidget(scroll)
        
        # 버튼
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("취소")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        if self.mode == "settings":
            save_settings_button = QPushButton("설정 저장")
            save_settings_button.clicked.connect(self.save_field_settings)
            button_layout.addWidget(save_settings_button)
        elif self.mode == "generate":
            create_button = QPushButton("생성")
            create_button.clicked.connect(self.create_from_controls)
            button_layout.addWidget(create_button)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
        self.setLayout(layout)

    def _migrate_branch_source_settings(self):
        selected_branches = self.settings.setdefault("branches", {})
        for setting in self.settings.get("fields", {}).values():
            field_type = setting.get("type")
            if field_type not in BRANCH_SOURCE_TYPES:
                continue
            if setting.get("branch_name") and not selected_branches.get(field_type):
                selected_branches[field_type] = setting["branch_name"]
            setting["type"] = "branch_value"
            setting["source_field"] = field_type
            setting["branch_value_key"] = "branch_name"
            setting.pop("branch_name", None)
            setting.pop("default_value", None)

    def _collect_controls(self, doc):
        controls = {}
        sdt_list = doc.element.findall('.//' + qn('w:sdt'))
        for idx, sdt in enumerate(sdt_list):
            info = self._content_control_info(sdt, idx)
            if not info:
                continue
            key = info["tag"] or info["alias"] or f"필드 {idx + 1}"
            control = controls.setdefault(key, {
                "tag": key,
                "alias": info["alias"],
                "placeholder": info["placeholder"],
                "current_text": info["current_text"],
                "indices": [],
            })
            control["indices"].append(idx)
            if not control["current_text"] and info["current_text"]:
                control["current_text"] = info["current_text"]
            if not control["alias"] and info["alias"]:
                control["alias"] = info["alias"]
            if not control["placeholder"] and info["placeholder"]:
                control["placeholder"] = info["placeholder"]
        return controls

    def _content_control_info(self, sdt, idx):
        sdt_pr = sdt.find(qn('w:sdtPr'))
        if sdt_pr is None:
            return None

        tag_elem = sdt_pr.find(qn('w:tag'))
        tag = tag_elem.get(qn('w:val')) if tag_elem is not None else f"필드 {idx + 1}"

        alias_elem = sdt_pr.find(qn('w:alias'))
        alias = alias_elem.get(qn('w:val')) if alias_elem is not None else ""

        placeholder_elem = sdt_pr.find(qn('w:placeholder'))
        placeholder = ""
        if placeholder_elem is not None:
            placeholder_text = placeholder_elem.find(qn('w:docPart'))
            if placeholder_text is not None:
                placeholder = placeholder_text.get(qn('w:val')) or ""

        current_text = self._sdt_text(sdt)
        return {
            "tag": tag,
            "alias": alias,
            "placeholder": placeholder,
            "current_text": current_text,
        }

    def _sdt_text(self, sdt):
        sdt_content = sdt.find(qn('w:sdtContent'))
        if sdt_content is None:
            return ""
        t_list = sdt_content.findall('.//' + qn('w:t'))
        return ''.join([t.text for t in t_list if t.text])

    def _add_field_row(self, form_layout, field_key, control, branch_member=False):
        setting = self.settings.setdefault("fields", {}).setdefault(field_key, {"type": "text"})
        if branch_member:
            setting["type"] = setting.get("field_type", "text")
        legacy_table_id = table_list_id(setting.get("type", ""))
        if legacy_table_id:
            setting["type"] = "table_list"
            setting.setdefault("table_list_id", legacy_table_id)

        display_name = field_key
        if len(control["indices"]) > 1:
            display_name = f"{field_key}({len(control['indices'])})"
        title_label = QLabel(f"<b>{display_name}</b>")
        description_texts = []
        if len(control["indices"]) > 1:
            description_texts.append(f"동일 이름 {len(control['indices'])}개 동시 수정")
        if control["alias"]:
            description_texts.append(f"별칭: {control['alias']}")
        if control["placeholder"]:
            description_texts.append(f"설명: {control['placeholder']}")
        description_label = QLabel(" | ".join(description_texts))
        description_label.setStyleSheet("color: gray; font-size: 10px;")
        description_label.setWordWrap(True)

        type_combo = QComboBox()
        type_combo.setMaximumWidth(220)
        for type_key, label in self._field_type_options():
            if branch_member and type_key == "branch":
                continue
            type_combo.addItem(label, type_key)
        type_index = type_combo.findData(setting.get("type", "text"))
        type_combo.setCurrentIndex(type_index if type_index >= 0 else 0)

        input_container = QWidget()
        input_layout = QHBoxLayout()
        input_layout.setContentsMargins(0, 0, 0, 0)
        input_container.setLayout(input_layout)

        self.field_widgets[field_key] = {
            "type_combo": type_combo,
            "input_container": input_container,
            "input_layout": input_layout,
            "setting": setting,
            "branch_member": branch_member,
        }
        self._refresh_type_combo_options(field_key)
        type_combo.currentIndexChanged.connect(lambda _, key=field_key: self._on_field_type_changed(key))

        if not branch_member:
            form_layout.addRow(title_label, None)
        form_layout.addRow("항목 형태" if branch_member else "형태", type_combo)
        form_layout.addRow("기본값", input_container)
        if description_texts and not branch_member:
            form_layout.addRow("", description_label)
        if not branch_member:
            form_layout.addRow("", QLabel(""))
        self._rebuild_field_input(field_key)

    def _add_branch_settings_group(self, form_layout, group_name, members):
        title_label = QLabel(f"<b>분기_{group_name}</b>")
        branch_type_combo = QComboBox()
        branch_type_combo.addItem(FIELD_TYPES["branch"], "branch")
        option_combo = QComboBox()
        stack = QStackedWidget()
        for member in members:
            option_combo.addItem(branch_control_parts(member)[1], member)
            page = QWidget()
            page_layout = QFormLayout(page)
            self._add_field_row(page_layout, member, self.controls[member], branch_member=True)
            stack.addWidget(page)
        option_combo.currentIndexChanged.connect(stack.setCurrentIndex)
        form_layout.addRow(title_label, None)
        form_layout.addRow("형태", branch_type_combo)
        form_layout.addRow("분기 항목", option_combo)
        form_layout.addRow("항목 설정", stack)
        form_layout.addRow("", QLabel(""))

    def _add_generate_rows(self, form_layout):
        configured_fields = self.settings.get("fields", {})
        if not configured_fields:
            form_layout.addRow(QLabel("설정된 콘텐츠 컨트롤 필드가 없습니다."), None)
            return

        rendered_branch_groups = set()
        for field_key, setting in configured_fields.items():
            if field_key not in self.controls:
                continue
            branch_parts = branch_control_parts(field_key) if setting.get("type") == "branch" else None
            if branch_parts:
                group_name, _option_name = branch_parts
                if group_name in rendered_branch_groups:
                    continue
                members = [
                    key for key, candidate in configured_fields.items()
                    if candidate.get("type") == "branch"
                    and branch_control_parts(key)
                    and branch_control_parts(key)[0] == group_name
                    and key in self.controls
                ]
                rendered_branch_groups.add(group_name)
                group_key = members[0]
                option_combo = QComboBox()
                option_combo.setMaximumWidth(220)
                stack = QStackedWidget()
                member_widgets = {}
                for member in members:
                    option_combo.addItem(branch_control_parts(member)[1], member)
                    member_setting = configured_fields[member]
                    member_type = member_setting.get("field_type", "text")
                    type_combo = QComboBox()
                    type_combo.addItem(self._field_type_label(member_type), member_type)
                    page = QWidget()
                    page_layout = QHBoxLayout(page)
                    page_layout.setContentsMargins(0, 0, 0, 0)
                    member_data = {
                        "type_combo": type_combo,
                        "input_container": page,
                        "input_layout": page_layout,
                        "setting": member_setting,
                        "branch_member": True,
                    }
                    self.field_widgets[member] = member_data
                    self._rebuild_field_input(member)
                    member_widgets[member] = member_data
                    stack.addWidget(page)
                option_combo.currentIndexChanged.connect(stack.setCurrentIndex)
                container = QWidget()
                container_layout = QVBoxLayout(container)
                container_layout.setContentsMargins(0, 0, 0, 0)
                select_row = QHBoxLayout()
                select_row.addWidget(option_combo)
                select_row.addStretch()
                container_layout.addLayout(select_row)
                container_layout.addWidget(stack)
                branch_type_combo = QComboBox()
                branch_type_combo.addItem(FIELD_TYPES["branch"], "branch")
                group_widgets = {
                    "type_combo": branch_type_combo,
                    "branch_members": members,
                    "branch_option_combo": option_combo,
                    "member_widgets": member_widgets,
                }
                form_layout.addRow(f"분기_{group_name}", container)
                self.generate_widgets[group_key] = group_widgets
                continue
            type_combo = QComboBox()
            field_type = setting.get("type", "text")
            type_combo.addItem(self._field_type_label(field_type), field_type)

            input_container = QWidget()
            input_layout = QHBoxLayout()
            input_layout.setContentsMargins(0, 0, 0, 0)
            input_container.setLayout(input_layout)

            self.field_widgets[field_key] = {
                "type_combo": type_combo,
                "input_container": input_container,
                "input_layout": input_layout,
                "setting": setting,
            }
            if branch_parts:
                self.field_widgets[field_key]["branch_members"] = members
            form_layout.addRow(field_key, input_container)
            self._rebuild_field_input(field_key)
            self.generate_widgets[field_key] = self.field_widgets[field_key]

        self._refresh_branch_combo_options()
        self._update_branch_value_fields()
        self._update_branch_summary()
        self._update_calculated_date_fields()

    def _default_value_for_generation(self, field_key, setting):
        field_type = setting.get("type", "text")
        control = self.controls.get(field_key, {})
        current_text = control.get("current_text", "")

        if field_type == "date":
            if setting.get("date_default_type", "today") == "field_calculation":
                return ""
            return QDate.currentDate().toString(setting.get("date_format", "yyyy-MM-dd"))
        if field_type == "amount":
            return setting.get("default_value") or ''.join(ch for ch in current_text if ch.isdigit())
        if field_type in {"branch_select", "branch_select_2"}:
            return setting.get("branch_name") or current_text
        if field_type == "branch_value":
            value = self._branch_value_from_setting(setting)
            if isinstance(value, dict):
                return value.get("path", "")
            return str(value)
        if field_type == "table_list" or table_list_id(field_type):
            return setting.get("default_value") or current_text
        if field_type == "folder":
            return setting.get("folder_path", "")
        return setting.get("default_value") or current_text

    def _field_type_options(self):
        return list(FIELD_TYPES.items())

    def _field_type_label(self, field_type):
        if table_list_id(field_type):
            return FIELD_TYPES["table_list"]
        for type_key, label in self._field_type_options():
            if type_key == field_type:
                return label
        return "삭제된 테이블 목록" if table_list_id(field_type) else FIELD_TYPES.get(field_type, "텍스트")

    def _table_list(self, field_type="", setting=None):
        target_id = (setting or {}).get("table_list_id") or table_list_id(field_type)
        return next((table for table in self.table_lists if table.get("id") == target_id), None)

    def _branch_value_from_setting(self, setting):
        source_type = self._source_type_from_setting(setting)
        branch = self._branch_from_source_setting(source_type)
        return self._branch_value(branch, setting.get("branch_value_key", "organization_code"))

    def _branch_from_source_setting(self, source_type):
        return self._selected_branch_from_source_type(source_type)

    def _on_field_type_changed(self, field_key):
        self._rebuild_field_input(field_key)
        self._refresh_type_combo_options()
        self._refresh_branch_combo_options(skip_field=field_key)
        self._refresh_branch_value_source_options(skip_field=field_key)
        self._update_branch_value_fields()
        self._update_branch_summary()

    def _refresh_type_combo_options(self, target_field_key=None):
        occupied_types = {
            widgets["type_combo"].currentData()
            for key, widgets in self.field_widgets.items()
            if key != target_field_key and widgets["type_combo"].currentData() in BRANCH_SOURCE_TYPES
        }
        targets = (
            [target_field_key]
            if target_field_key
            else list(self.field_widgets.keys())
        )
        for field_key in targets:
            widgets = self.field_widgets.get(field_key)
            if not widgets:
                continue
            type_combo = widgets["type_combo"]
            current_type = type_combo.currentData() or widgets["setting"].get("type", "text")
            if table_list_id(current_type):
                current_type = "table_list"
            type_combo.blockSignals(True)
            type_combo.clear()
            for type_key, label in self._field_type_options():
                if widgets.get("branch_member") and type_key == "branch":
                    continue
                if type_key in BRANCH_SOURCE_TYPES and type_key in occupied_types and type_key != current_type:
                    continue
                type_combo.addItem(label, type_key)
            type_index = type_combo.findData(current_type)
            type_combo.setCurrentIndex(type_index if type_index >= 0 else 0)
            type_combo.blockSignals(False)

    def _refresh_branch_value_source_options(self, skip_field=None):
        for field_key, widgets in self.field_widgets.items():
            if field_key == skip_field or widgets["type_combo"].currentData() != "branch_value":
                continue
            source_combo = widgets.get("source_combo")
            if not source_combo:
                continue
            current_source = source_combo.currentData()
            source_combo.blockSignals(True)
            source_combo.clear()
            for source_type, label in self._branch_source_options(field_key):
                source_combo.addItem(label, source_type)
            source_index = source_combo.findData(current_source)
            if source_index < 0:
                source_index = source_combo.findData(widgets["setting"].get("source_field", "branch_select"))
            source_combo.setCurrentIndex(source_index if source_index >= 0 else 0)
            source_combo.blockSignals(False)
            self._update_branch_value_field(field_key)

    def _on_global_branch_changed(self, _index=None):
        self._refresh_branch_value_source_options()
        self._update_branch_value_fields()

    def _clear_layout(self, layout):
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()

    def _rebuild_field_input(self, field_key):
        widgets = self.field_widgets[field_key]
        layout = widgets["input_layout"]
        self._clear_layout(layout)

        field_type = widgets["type_combo"].currentData() or "text"
        control = self.controls[field_key]
        setting = widgets["setting"]
        setting["type"] = field_type

        if field_type == "branch":
            parts = branch_control_parts(field_key)
            branch_members = widgets.get("branch_members", [field_key])
            if self.mode == "settings":
                hint = QLabel(
                    f"{parts[0]} 그룹 / {parts[1]} 항목" if parts else
                    "이름을 분기_필드명_내용명 형식으로 지정해 주세요."
                )
                hint.setStyleSheet("color: gray;")
                layout.addWidget(hint)
                widgets["value_widget"] = hint
                return

            option_combo = QComboBox()
            option_combo.setMaximumWidth(220)
            for member in branch_members:
                member_parts = branch_control_parts(member)
                if member_parts:
                    option_combo.addItem(member_parts[1], member)
            value_edit = QLineEdit()
            value_edit.setPlaceholderText("선택한 칸에 입력할 내용")
            layout.addWidget(option_combo)
            layout.addWidget(value_edit)
            widgets["branch_option_combo"] = option_combo
            widgets["value_widget"] = value_edit
            return

        if field_type == "date":
            date_edit = QDateEdit()
            date_edit.setCalendarPopup(True)
            date_edit.setDisplayFormat("yyyy-MM-dd")
            date_default_type = setting.get("date_default_type", "today")
            if date_default_type == "field_calculation":
                date_edit.setDate(QDate.currentDate())
            else:
                date_edit.setDate(QDate.currentDate())

            format_combo = QComboBox()
            format_combo.setMaximumWidth(170)
            for value, label in DATE_FORMATS.items():
                format_combo.addItem(label, value)
            format_combo.setCurrentText(setting.get("date_format", "yyyy-MM-dd"))
            index = format_combo.findData(setting.get("date_format", "yyyy-MM-dd"))
            format_combo.setCurrentIndex(index if index >= 0 else 0)

            if self.mode == "settings":
                default_type_combo = QComboBox()
                default_type_combo.setMaximumWidth(150)
                default_type_combo.addItem("오늘", "today")
                default_type_combo.addItem("특정 필드 계산", "field_calculation")
                type_index = default_type_combo.findData(date_default_type)
                default_type_combo.setCurrentIndex(type_index if type_index >= 0 else 0)

                source_combo = QComboBox()
                source_combo.setMaximumWidth(200)
                for source_key, source_setting in self.settings.get("fields", {}).items():
                    if source_key != field_key and source_setting.get("type") == "date":
                        source_combo.addItem(source_key, source_key)
                source_index = source_combo.findData(setting.get("source_date_field", ""))
                source_combo.setCurrentIndex(source_index if source_index >= 0 else 0)

                sign_combo = QComboBox()
                sign_combo.addItem("+", 1)
                sign_combo.addItem("-", -1)
                sign_index = sign_combo.findData(int(setting.get("date_offset_sign", 1) or 1))
                sign_combo.setCurrentIndex(sign_index if sign_index >= 0 else 0)

                days_edit = QLineEdit(str(setting.get("date_offset_days", 0) or 0))
                days_edit.setValidator(QRegularExpressionValidator(QRegularExpression(r"\d{0,5}")))
                days_edit.setMaximumWidth(65)

                date_options = QWidget()
                date_options_layout = QVBoxLayout()
                date_options_layout.setContentsMargins(0, 0, 0, 0)
                first_row = QHBoxLayout()
                first_row.addWidget(QLabel("날짜"))
                first_row.addWidget(date_edit)
                first_row.addWidget(QLabel("표시 형식"))
                first_row.addWidget(format_combo)
                first_row.addWidget(QLabel("기준"))
                first_row.addWidget(default_type_combo)
                first_row.addStretch()
                calculation_row = QHBoxLayout()
                calculation_label = QLabel("계산")
                calculation_row.addWidget(calculation_label)
                calculation_row.addWidget(source_combo)
                calculation_row.addWidget(sign_combo)
                calculation_row.addWidget(days_edit)
                days_label = QLabel("일")
                calculation_row.addWidget(days_label)
                calculation_row.addStretch()
                date_options_layout.addLayout(first_row)
                date_options_layout.addLayout(calculation_row)
                date_options.setLayout(date_options_layout)
                layout.addWidget(date_options)

                def update_calculation_controls():
                    calculated = default_type_combo.currentData() == "field_calculation"
                    date_edit.setReadOnly(calculated)
                    source_combo.setVisible(calculated)
                    sign_combo.setVisible(calculated)
                    days_edit.setVisible(calculated)
                    calculation_label.setVisible(calculated)
                    days_label.setVisible(calculated)

                default_type_combo.currentIndexChanged.connect(lambda _: update_calculation_controls())
                default_type_combo.currentIndexChanged.connect(lambda _: self._update_calculated_date_fields())
                source_combo.currentIndexChanged.connect(lambda _: self._update_calculated_date_fields())
                sign_combo.currentIndexChanged.connect(lambda _: self._update_calculated_date_fields())
                days_edit.textChanged.connect(lambda _: self._update_calculated_date_fields())
                update_calculation_controls()
                widgets["date_default_type_combo"] = default_type_combo
                widgets["source_date_combo"] = source_combo
                widgets["date_offset_sign_combo"] = sign_combo
                widgets["date_offset_days_edit"] = days_edit
            elif date_default_type == "field_calculation":
                date_edit.setReadOnly(True)
                layout.addWidget(date_edit)
            else:
                layout.addWidget(date_edit)
            widgets["value_widget"] = date_edit
            widgets["format_combo"] = format_combo
            date_edit.dateChanged.connect(lambda _: self._update_calculated_date_fields())
            return

        if field_type == "amount":
            amount_edit = QLineEdit()
            amount_edit.setMaximumWidth(220)
            amount_edit.setValidator(QRegularExpressionValidator(QRegularExpression(r"\d*")))
            amount_edit.setText(''.join(ch for ch in str(setting.get("default_value") or control["current_text"]) if ch.isdigit()))
            comma_check = QCheckBox("1000단위 쉼표")
            comma_check.setChecked(bool(setting.get("use_comma", True)))
            layout.addWidget(amount_edit)
            if self.mode == "settings":
                layout.addWidget(comma_check)
            widgets["value_widget"] = amount_edit
            widgets["comma_check"] = comma_check
            return

        if field_type in {"branch_select", "branch_select_2"}:
            branch_combo = QComboBox()
            branch_combo.setMaximumWidth(240)
            self._populate_branch_combo(field_key, branch_combo, setting.get("branch_name") or control["current_text"])
            branch_combo.currentIndexChanged.connect(lambda _, key=field_key: self._on_branch_selection_changed(key))
            layout.addWidget(branch_combo)
            widgets["value_widget"] = branch_combo
            return

        if field_type == "branch_value":
            if self.mode != "settings":
                preview = QLineEdit()
                preview.setReadOnly(True)
                layout.addWidget(preview)
                widgets["value_widget"] = preview
                self._update_branch_value_field(field_key)
                return

            source_combo = QComboBox()
            source_combo.setMaximumWidth(180)
            for source_type, label in self._branch_source_options(field_key):
                source_combo.addItem(label, source_type)
            source_index = source_combo.findData(self._source_type_from_setting(setting))
            source_combo.setCurrentIndex(source_index if source_index >= 0 else 0)

            value_combo = QComboBox()
            value_combo.setMaximumWidth(220)
            for value_key, label in self._branch_value_field_options().items():
                value_combo.addItem(label, value_key)
            value_index = value_combo.findData(setting.get("branch_value_key", "organization_code"))
            value_combo.setCurrentIndex(value_index if value_index >= 0 else 0)

            preview = QLineEdit()
            preview.setReadOnly(True)
            source_combo.currentIndexChanged.connect(lambda _: self._update_branch_value_field(field_key))
            value_combo.currentIndexChanged.connect(lambda _: self._update_branch_value_field(field_key))

            layout.addWidget(source_combo)
            layout.addWidget(value_combo)
            layout.addWidget(preview)
            widgets["source_combo"] = source_combo
            widgets["value_key_combo"] = value_combo
            widgets["value_widget"] = preview
            self._update_branch_value_field(field_key)
            return

        if field_type == "table_list" or table_list_id(field_type):
            legacy_table_id = table_list_id(field_type)
            if legacy_table_id and not setting.get("table_list_id"):
                setting["table_list_id"] = legacy_table_id
            setting["type"] = "table_list"

            table_combo = QComboBox()
            table_combo.setMaximumWidth(220)
            for table in self.table_lists:
                if table.get("id"):
                    table_combo.addItem(str(table.get("name", "테이블 목록")), table.get("id"))

            value_combo = QComboBox()
            value_combo.setMaximumWidth(260)
            preferred_table_id = setting.get("table_list_id", "")
            table_index = table_combo.findData(preferred_table_id)
            if table_index >= 0:
                table_combo.setCurrentIndex(table_index)

            def populate_values():
                preferred = setting.get("default_value") or control["current_text"]
                value_combo.clear()
                table = next(
                    (item for item in self.table_lists if item.get("id") == table_combo.currentData()),
                    None,
                )
                for value in table_list_options(table or {}):
                    value_combo.addItem(value, value)
                preferred_index = value_combo.findData(preferred)
                if preferred_index >= 0:
                    value_combo.setCurrentIndex(preferred_index)

            populate_values()
            table_combo.currentIndexChanged.connect(lambda _: populate_values())
            layout.addWidget(table_combo)
            layout.addWidget(value_combo)
            widgets["table_combo"] = table_combo
            widgets["value_widget"] = value_combo
            return

        if field_type == "folder":
            folder_edit = QLineEdit(setting.get("folder_path", ""))
            folder_edit.setReadOnly(True)

            if self.mode == "settings":
                select_button = QPushButton("폴더 선택")

                def select_folder():
                    selected = QFileDialog.getExistingDirectory(
                        self,
                        "이미지 폴더 선택",
                        folder_edit.text() or str(self.root),
                    )
                    if selected:
                        folder_edit.setText(selected)

                select_button.clicked.connect(select_folder)
                source_combo = QComboBox()
                source_combo.setMaximumWidth(220)
                for source_key in self.controls:
                    if source_key != field_key:
                        source_combo.addItem(source_key, source_key)
                source_index = source_combo.findData(setting.get("value_field", ""))
                source_combo.setCurrentIndex(source_index if source_index >= 0 else 0)

                size_validator = QRegularExpressionValidator(QRegularExpression(r"[1-9]\d{0,2}"))
                width_edit = QLineEdit(str(setting.get("width", 30) or 30))
                width_edit.setValidator(size_validator)
                width_edit.setMaximumWidth(55)
                height_edit = QLineEdit(str(setting.get("height", 30) or 30))
                height_edit.setValidator(QRegularExpressionValidator(QRegularExpression(r"[1-9]\d{0,2}")))
                height_edit.setMaximumWidth(55)

                folder_options = QWidget()
                folder_options_layout = QVBoxLayout()
                folder_options_layout.setContentsMargins(0, 0, 0, 0)
                folder_row = QHBoxLayout()
                folder_row.addWidget(folder_edit)
                folder_row.addWidget(select_button)
                setting_row = QHBoxLayout()
                setting_row.addWidget(QLabel("값필드"))
                setting_row.addWidget(source_combo)
                setting_row.addWidget(QLabel("너비(mm)"))
                setting_row.addWidget(width_edit)
                setting_row.addWidget(QLabel("높이(mm)"))
                setting_row.addWidget(height_edit)
                setting_row.addStretch()
                folder_options_layout.addLayout(folder_row)
                folder_options_layout.addLayout(setting_row)
                folder_options.setLayout(folder_options_layout)
                layout.addWidget(folder_options)
                widgets["source_field_combo"] = source_combo
                widgets["width_edit"] = width_edit
                widgets["height_edit"] = height_edit
            else:
                source_name = setting.get("value_field", "")
                width = setting.get("width", 30) or 30
                height = setting.get("height", 30) or 30
                folder_edit.setText(
                    f"{setting.get('folder_path', '')}  ←  {source_name}  ({width}×{height}mm)"
                )
                layout.addWidget(folder_edit)
            widgets["value_widget"] = folder_edit
            return

        text_edit = QLineEdit()
        text_edit.setText(setting.get("default_value") or control["current_text"])
        text_edit.setPlaceholderText(control["placeholder"] or control["alias"] or "입력하세요")
        layout.addWidget(text_edit)
        widgets["value_widget"] = text_edit

    def _parse_date(self, text):
        parsed = self._parse_date_or_none(text)
        return parsed if parsed else QDate.currentDate()

    def _parse_date_or_none(self, text):
        digits = ''.join(ch if ch.isdigit() else '-' for ch in str(text or "")).strip('-')
        parts = [part for part in digits.split('-') if part]
        if len(parts) >= 3:
            date = QDate(int(parts[0]), int(parts[1]), int(parts[2]))
            if date.isValid():
                return date
        return None

    def _date_field_value(self, field_key, visited=None):
        visited = set(visited or ())
        if field_key in visited:
            return ""
        visited.add(field_key)
        widgets = self.field_widgets.get(field_key)
        if not widgets:
            return ""
        setting = widgets.get("setting", {})
        default_type_combo = widgets.get("date_default_type_combo")
        default_type = (
            default_type_combo.currentData()
            if default_type_combo
            else setting.get("date_default_type", "today")
        )
        date_format = widgets["format_combo"].currentData()
        if default_type != "field_calculation":
            return widgets["value_widget"].date().toString(date_format)

        source_combo = widgets.get("source_date_combo")
        source_key = source_combo.currentData() if source_combo else setting.get("source_date_field", "")
        source_text = self._date_field_value(source_key, visited)
        source_date = self._parse_date_or_none(source_text)
        if not source_date:
            return ""
        sign_combo = widgets.get("date_offset_sign_combo")
        days_edit = widgets.get("date_offset_days_edit")
        sign = sign_combo.currentData() if sign_combo else int(setting.get("date_offset_sign", 1) or 1)
        days = int(days_edit.text() or 0) if days_edit else int(setting.get("date_offset_days", 0) or 0)
        return source_date.addDays(sign * days).toString(date_format)

    def _update_calculated_date_fields(self):
        for field_key, widgets in self.field_widgets.items():
            if widgets["type_combo"].currentData() != "date":
                continue
            setting = widgets.get("setting", {})
            default_type_combo = widgets.get("date_default_type_combo")
            default_type = (
                default_type_combo.currentData()
                if default_type_combo
                else setting.get("date_default_type", "today")
            )
            if default_type != "field_calculation":
                continue
            calculated = self._parse_date_or_none(self._date_field_value(field_key))
            if calculated:
                date_edit = widgets["value_widget"]
                date_edit.blockSignals(True)
                date_edit.setDate(calculated)
                date_edit.blockSignals(False)

    def _branch_key(self, branch):
        if not branch:
            return ""
        return str(branch.get("organization_code") or branch.get("branch_name") or "")

    def _branch_name(self, branch):
        return branch.get("branch_name", "") if branch else ""

    def _branch_combo_current_key(self, combo):
        return self._branch_key(combo.currentData()) if combo else ""

    def _selected_branch_keys(self, exclude_field_key=None):
        selected = set()
        for field_key, widgets in self.field_widgets.items():
            if field_key == exclude_field_key:
                continue
            if widgets["type_combo"].currentData() not in BRANCH_SOURCE_TYPES:
                continue
            key = self._branch_combo_current_key(widgets.get("value_widget"))
            if key:
                selected.add(key)
        return selected

    def _populate_branch_combo(self, field_key, combo, preferred_branch_name=None):
        preferred_branch_name = preferred_branch_name or self._branch_name(combo.currentData())
        excluded_keys = self._selected_branch_keys(exclude_field_key=field_key)

        combo.blockSignals(True)
        combo.clear()
        for branch in self.branches:
            branch_key = self._branch_key(branch)
            if branch_key in excluded_keys:
                continue
            combo.addItem(branch.get("branch_name", ""), branch)
        self._select_branch_combo(combo, preferred_branch_name)
        combo.blockSignals(False)

    def _branch_value_field_options(self):
        options = dict(STANDARD_BRANCH_VALUE_FIELDS)
        for field in self.branch_fields:
            key = field.get("key")
            label = field.get("label", key)
            if key and label:
                options[key] = label
        return options

    def _branch_field_definition(self, key):
        for field in self.branch_fields:
            if field.get("key") == key:
                return field
        return None

    def _branch_value(self, branch, key):
        if not branch:
            return ""
        if key in STANDARD_BRANCH_VALUE_FIELDS:
            return str(branch.get(key, ""))

        field = self._branch_field_definition(key)
        custom_value = branch.get("custom_fields", {}).get(key, {})
        if not isinstance(custom_value, dict):
            return str(custom_value)
        if field and field.get("type") == "image":
            path_text = custom_value.get("path", "")
            if not path_text:
                return {"type": "image", "path": "", "width": 30, "height": 30}
            path = Path(path_text)
            if path_text and not path.is_absolute():
                path = self.root / path_text
            return {
                "type": "image",
                "path": str(path),
                "width": int(custom_value.get("width", 30) or 30),
                "height": int(custom_value.get("height", 30) or 30),
            }
        return str(custom_value.get("value", ""))

    def _select_branch_combo(self, combo, branch_name):
        for index in range(combo.count()):
            branch = combo.itemData(index)
            if branch and branch.get("branch_name") == branch_name:
                combo.setCurrentIndex(index)
                return

    def _refresh_branch_combo_options(self, skip_field=None):
        for field_key, widgets in self.field_widgets.items():
            if field_key == skip_field or widgets["type_combo"].currentData() not in BRANCH_SOURCE_TYPES:
                continue
            combo = widgets.get("value_widget")
            if not isinstance(combo, QComboBox):
                continue
            self._populate_branch_combo(field_key, combo, self._branch_name(combo.currentData()))
            branch = combo.currentData()
            if self.mode == "settings":
                widgets["setting"]["branch_name"] = self._branch_name(branch)

    def _on_branch_selection_changed(self, field_key):
        widgets = self.field_widgets.get(field_key)
        if not widgets:
            return
        combo = widgets.get("value_widget")
        branch = combo.currentData() if isinstance(combo, QComboBox) else None
        if self.mode == "settings":
            widgets["setting"]["branch_name"] = self._branch_name(branch)
        self._refresh_branch_combo_options(skip_field=field_key)
        self._refresh_branch_value_source_options()
        self._update_branch_value_fields()
        self._update_branch_summary()

    def _update_branch_summary(self):
        if not hasattr(self, "branch_summary_label"):
            return
        parts = []
        for source_type, label in BRANCH_SOURCE_TYPES.items():
            branch = self._selected_branch_from_source_type(source_type)
            if branch:
                parts.append(f"{label}: {self._branch_name(branch)}")
        self.branch_summary_label.setText(" | ".join(parts) if parts else "선택된 지점 없음")

    def _branch_source_options(self, current_field_key=None):
        options = []
        for source_type, label in BRANCH_SOURCE_TYPES.items():
            branch = self._selected_branch_from_source_type(source_type)
            branch_name = self._branch_name(branch)
            option_label = f"{branch_name} ({label})" if branch_name else label
            options.append((source_type, option_label))
        return options

    def _source_type_from_setting(self, setting):
        source = setting.get("source_field", "branch_select")
        if source in BRANCH_SOURCE_TYPES:
            return source
        widgets = self.field_widgets.get(source)
        if widgets and widgets["type_combo"].currentData() in BRANCH_SOURCE_TYPES:
            return widgets["type_combo"].currentData()
        return "branch_select"

    def _selected_branch_from_source_type(self, source_type):
        combo = self.branch_combos.get(source_type)
        return combo.currentData() if combo else None

    def _update_branch_value_fields(self):
        for field_key, widgets in self.field_widgets.items():
            if widgets["type_combo"].currentData() == "branch_value":
                self._update_branch_value_field(field_key)

    def _update_branch_value_field(self, field_key):
        widgets = self.field_widgets.get(field_key, {})
        source_combo = widgets.get("source_combo")
        value_key_combo = widgets.get("value_key_combo")
        preview = widgets.get("value_widget")
        if not preview:
            return
        source_type = source_combo.currentData() if source_combo else self._source_type_from_setting(widgets.get("setting", {}))
        value_key = value_key_combo.currentData() if value_key_combo else widgets.get("setting", {}).get("branch_value_key", "organization_code")
        branch = self._selected_branch_from_source_type(source_type)
        value = ""
        if branch:
            value = self._branch_value(branch, value_key)
            if isinstance(value, dict):
                value = value.get("path", "")
        preview.setText(value)

    def _field_value(self, field_key):
        widgets = self.field_widgets[field_key]
        field_type = widgets["type_combo"].currentData() or "text"
        value_widget = widgets.get("value_widget")

        if field_type == "date":
            return self._date_field_value(field_key)
        if field_type == "amount":
            text = value_widget.text()
            if widgets["comma_check"].isChecked() and text:
                return f"{int(text):,}"
            return text
        if field_type in {"branch_select", "branch_select_2"}:
            branch = value_widget.currentData()
            return branch.get("branch_name", "") if branch else ""
        if field_type == "branch_value":
            source_combo = widgets.get("source_combo")
            value_key_combo = widgets.get("value_key_combo")
            source_type = source_combo.currentData() if source_combo else self._source_type_from_setting(widgets.get("setting", {}))
            value_key = value_key_combo.currentData() if value_key_combo else widgets.get("setting", {}).get("branch_value_key", "organization_code")
            branch = self._selected_branch_from_source_type(source_type)
            return self._branch_value(branch, value_key)
        if field_type == "table_list" or table_list_id(field_type):
            return value_widget.currentData() or "" if value_widget else ""
        if field_type == "folder":
            setting = widgets.get("setting", {})
            source_key = setting.get("value_field", "")
            source_value = self._field_value(source_key) if source_key in self.field_widgets else ""
            return self._folder_image_value(setting, source_value)
        return value_widget.text() if value_widget else ""

    def _generation_values(self):
        values = {}
        for field_key, widgets in self.generate_widgets.items():
            if widgets["type_combo"].currentData() != "branch":
                values[field_key] = self._field_value(field_key)
                continue
            selected = widgets.get("branch_option_combo").currentData()
            if selected:
                values[selected] = self._field_value(selected)
        return values

    def _validate_generation_inputs(self):
        errors = []
        for field_key, widgets in self.generate_widgets.items():
            field_type = widgets["type_combo"].currentData() or "text"
            value_widget = widgets.get("value_widget")
            if field_type == "amount":
                text = value_widget.text().strip() if value_widget else ""
                if text and not text.isdigit():
                    errors.append(f"{field_key}: 숫자만 입력할 수 있습니다.")
            elif field_type == "date":
                if not value_widget or not value_widget.date().isValid():
                    errors.append(f"{field_key}: 올바른 날짜를 입력해 주세요.")
            elif field_type in {"branch_select", "branch_select_2"}:
                if not value_widget or value_widget.currentData() is None:
                    errors.append(f"{field_key}: 지점을 선택해 주세요.")
            elif field_type == "table_list" or table_list_id(field_type):
                if not value_widget or value_widget.currentData() is None:
                    errors.append(f"{field_key}: 테이블 값을 선택해 주세요.")
        return errors

    def _collect_field_settings(self):
        fields = {}
        for field_key, widgets in self.field_widgets.items():
            field_type = widgets["type_combo"].currentData() or "text"
            setting = {"type": field_type}
            if field_type == "date":
                setting["date_format"] = widgets["format_combo"].currentData()
                setting["date_default_type"] = widgets["date_default_type_combo"].currentData()
                if setting["date_default_type"] == "field_calculation":
                    setting["source_date_field"] = widgets["source_date_combo"].currentData() or ""
                    setting["date_offset_sign"] = widgets["date_offset_sign_combo"].currentData() or 1
                    setting["date_offset_days"] = int(widgets["date_offset_days_edit"].text() or 0)
                else:
                    setting.pop("source_date_field", None)
                    setting.pop("date_offset_sign", None)
                    setting.pop("date_offset_days", None)
                setting.pop("default_value", None)
            elif field_type == "amount":
                setting["use_comma"] = widgets["comma_check"].isChecked()
                setting["default_value"] = widgets["value_widget"].text()
            elif field_type in {"branch_select", "branch_select_2"}:
                branch = widgets["value_widget"].currentData()
                setting["branch_name"] = branch.get("branch_name", "") if branch else ""
            elif field_type == "branch_value":
                setting["source_field"] = widgets["source_combo"].currentData()
                setting["branch_value_key"] = widgets["value_key_combo"].currentData()
            elif field_type == "table_list" or table_list_id(field_type):
                setting["type"] = "table_list"
                setting["table_list_id"] = widgets["table_combo"].currentData() or ""
                setting["default_value"] = widgets["value_widget"].currentData() or ""
            elif field_type == "folder":
                setting["folder_path"] = widgets["value_widget"].text().strip()
                setting["value_field"] = widgets["source_field_combo"].currentData() or ""
                setting["width"] = int(widgets["width_edit"].text() or 0)
                setting["height"] = int(widgets["height_edit"].text() or 0)
            elif field_type == "branch":
                setting = {"type": "branch"}
            else:
                value_widget = widgets.get("value_widget")
                setting["default_value"] = value_widget.text() if value_widget else ""
            if widgets.get("branch_member"):
                setting["field_type"] = setting.pop("type")
                setting["type"] = "branch"
            fields[field_key] = setting
        # Selecting 분기 on one member configures every matching member together.
        branch_groups = {
            branch_control_parts(key)[0]
            for key, setting in fields.items()
            if setting.get("type") == "branch" and branch_control_parts(key)
        }
        for key in self.controls:
            parts = branch_control_parts(key)
            if parts and parts[0] in branch_groups:
                fields.setdefault(key, {"type": "branch", "field_type": "text"})
        return {
            "version": 1,
            "branches": {
                source_type: self._branch_name(combo.currentData())
                for source_type, combo in self.branch_combos.items()
            },
            "fields": fields,
        }

    def save_field_settings(self):
        settings = self._collect_field_settings()
        errors = []
        for field_key, setting in settings.get("fields", {}).items():
            if setting.get("type") == "branch" and not branch_control_parts(field_key):
                errors.append(f"{field_key}: 분기_필드명_내용명 형식의 필드 이름이 필요합니다.")
            if setting.get("type") == "date" and setting.get("date_default_type") == "field_calculation":
                source_key = setting.get("source_date_field", "")
                if not source_key or settings["fields"].get(source_key, {}).get("type") != "date":
                    errors.append(f"{field_key}: 계산 기준이 될 날짜 필드를 선택해 주세요.")
            if setting.get("type") != "folder":
                continue
            folder_path = setting.get("folder_path", "")
            source_key = setting.get("value_field", "")
            if not folder_path or not Path(folder_path).is_dir():
                errors.append(f"{field_key}: 사용할 이미지 폴더를 선택해 주세요.")
            if not source_key:
                errors.append(f"{field_key}: 파일명과 비교할 값필드를 선택해 주세요.")
            elif settings["fields"].get(source_key, {}).get("type") == "folder":
                errors.append(f"{field_key}: 폴더의 이미지 형태가 아닌 값필드를 선택해 주세요.")
            if setting.get("width", 0) <= 0 or setting.get("height", 0) <= 0:
                errors.append(f"{field_key}: 이미지 너비와 높이를 1~999mm로 입력해 주세요.")
        for field_key, setting in settings.get("fields", {}).items():
            if setting.get("type") != "date" or setting.get("date_default_type") != "field_calculation":
                continue
            visited = {field_key}
            source_key = setting.get("source_date_field", "")
            while source_key:
                if source_key in visited:
                    errors.append(f"{field_key}: 날짜 계산 필드가 서로 순환 참조하고 있습니다.")
                    break
                visited.add(source_key)
                source_setting = settings["fields"].get(source_key, {})
                if source_setting.get("date_default_type") != "field_calculation":
                    break
                source_key = source_setting.get("source_date_field", "")
        if errors:
            QMessageBox.warning(self, "필드 설정 확인", "\n".join(errors))
            return

        self.settings = settings
        self.all_settings.setdefault("documents", {})[self.file_path.name] = self.settings
        write_json_file(self.settings_path, self.all_settings)
        QMessageBox.information(self, "저장 완료", "필드 설정을 저장했습니다.")
    
    def _load_docx2pdf(self):
        try:
            return importlib.import_module('docx2pdf')
        except ImportError:
            return None

    def _install_docx2pdf(self):
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'docx2pdf'])
            importlib.invalidate_caches()
            return importlib.import_module('docx2pdf')
        except Exception as e:
            raise RuntimeError(f"docx2pdf 설치 실패: {e}") from e

    def create_from_controls(self):
        """입력한 값으로 새 문서를 생성(PDF 변환). 원본은 변경하지 않음."""
        log("DocumentTool", f"문서 생성 시도: {self.file_path.name}", level="INFO")
        try:
            title = self.title_edit.text().strip() if self.title_edit else self.file_path.stem
            if not title:
                QMessageBox.warning(self, "제목 필요", "저장 제목을 입력해 주세요.")
                return

            validation_errors = self._validate_generation_inputs()
            if validation_errors:
                QMessageBox.warning(self, "입력 확인", "\n".join(validation_errors))
                return

            values_dict = self._generation_values()

            if not values_dict:
                QMessageBox.warning(self, "설정 필요", "생성할 콘텐츠 컨트롤 필드가 없습니다.")
                return

            progress = self._show_progress("문서 생성", "문서 값을 적용하고 있습니다...", 4)

            def update_progress(step, message):
                progress.setLabelText(message)
                progress.setValue(step)
                QApplication.processEvents()

            try:
                out_pdf = self._create_pdf(
                    values_dict,
                    output_title=title,
                    ask_on_duplicate=True,
                    status_callback=update_progress,
                )
            finally:
                progress.close()
            if out_pdf is None:
                return
            QMessageBox.information(self, '완료', f'PDF 생성 완료:\n{out_pdf}\n\n')

            self.accept()
        except Exception as e:
            log("DocumentTool", f"문서 생성 오류: {e}", level="ERROR")
            QMessageBox.critical(self, "오류", f"생성 중 오류: {e}")

    def create_documents_from_rows(self, rows):
        outputs = []
        progress = self._show_progress("문서 일괄 생성", f"PDF 문서를 생성하고 있습니다... (0/{len(rows)})", len(rows))
        try:
            for index, values_dict in enumerate(rows, 1):
                values_dict = self._prepare_excel_row_values(values_dict, row_number=index + 1)
                self._validate_row_values(values_dict, row_number=index + 1)
                output_title = f"{self.file_path.stem}_filled_{index:03d}"

                def update_batch_status(_step, message, current=index):
                    progress.setLabelText(f"{message} ({current}/{len(rows)})")
                    QApplication.processEvents()

                outputs.append(
                    self._create_pdf(
                        values_dict,
                        output_title=output_title,
                        ask_on_duplicate=False,
                        status_callback=update_batch_status,
                    )
                )
                progress.setLabelText(f"PDF 문서를 생성하고 있습니다... ({index}/{len(rows)})")
                progress.setValue(index)
                QApplication.processEvents()
        finally:
            progress.close()
        return outputs

    def _prepare_excel_row_values(self, values_dict, row_number):
        """엑셀의 이미지 필드 경로를 문서 삽입용 이미지 값으로 변환한다."""
        prepared = dict(values_dict)
        for field_key, setting in self.settings.get("fields", {}).items():
            if not self._is_image_field_setting(setting):
                continue

            path_text = str(prepared.get(field_key, "") or "").strip().strip('"')
            if not path_text:
                prepared[field_key] = ""
                continue

            image_path = Path(path_text)
            if not image_path.is_absolute():
                image_path = self.root / image_path
            if not image_path.is_file():
                raise ValueError(f"{row_number}행 {field_key}: 이미지 파일을 찾을 수 없습니다.\n{image_path}")
            if image_path.suffix.casefold() not in IMAGE_EXTENSIONS:
                raise ValueError(f"{row_number}행 {field_key}: 지원하지 않는 이미지 형식입니다.\n{image_path}")

            width, height = self._image_size_for_field(field_key)
            prepared[field_key] = {
                "type": "image",
                "path": str(image_path),
                "width": width,
                "height": height,
            }
        return prepared

    def _is_image_field_setting(self, setting):
        if setting.get("type") == "image":
            return True
        if setting.get("type") != "branch_value":
            return False
        field = self._branch_field_definition(setting.get("branch_value_key", ""))
        return bool(field and field.get("type") == "image")

    def _image_size_for_field(self, field_key):
        """선택된 지점의 이미지 규격을 재사용하고, 없으면 30mm를 사용한다."""
        if field_key in self.field_widgets:
            current_value = self._field_value(field_key)
            if isinstance(current_value, dict) and current_value.get("type") == "image":
                return (
                    int(current_value.get("width", 30) or 30),
                    int(current_value.get("height", 30) or 30),
                )
        return 30, 30

    def _show_progress(self, title, message, maximum=0):
        progress = QProgressDialog(message, "", 0, maximum, self)
        progress.setWindowTitle(title)
        progress.setCancelButton(None)
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(0)
        progress.setAutoClose(False)
        progress.setAutoReset(False)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()
        return progress

    def _create_pdf(self, values_dict, output_title=None, ask_on_duplicate=False, status_callback=None):
        def report(step, message):
            if status_callback:
                status_callback(step, message)

        report(1, "문서 값을 적용하고 있습니다...")
        new_doc = Document(self.file_path)
        self._apply_values_to_doc(new_doc, values_dict)

        maked_folder = self.file_path.parent / "maked"
        maked_folder.mkdir(exist_ok=True)
        out_pdf = self._resolve_output_pdf_path(maked_folder, output_title or f"{self.file_path.stem}_filled", ask_on_duplicate)
        if out_pdf is None:
            return None

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            temp_docx_path = Path(tmp_file.name)
        try:
            report(2, "변환할 Word 문서를 준비하고 있습니다...")
            new_doc.save(str(temp_docx_path))
            docx2pdf = self._ensure_docx2pdf()
            report(3, "Word에서 PDF로 변환하고 있습니다...")
            docx2pdf.convert(str(temp_docx_path), str(out_pdf))
            report(4, "PDF 저장을 완료하고 있습니다...")
            log("DocumentTool", f"PDF 생성 완료: {out_pdf}", level="INFO")
            return out_pdf
        finally:
            try:
                if temp_docx_path.exists():
                    temp_docx_path.unlink()
            except Exception:
                pass

    def _validate_row_values(self, values_dict, row_number):
        errors = []
        for field_key, setting in self.settings.get("fields", {}).items():
            value = str(values_dict.get(field_key, "")).strip()
            field_type = setting.get("type", "text")
            if field_type == "amount" and value and not value.replace(",", "").isdigit():
                errors.append(f"{row_number}행 {field_key}: 숫자만 입력할 수 있습니다.")
            elif field_type == "date" and value and not self._parse_date_or_none(value):
                errors.append(f"{row_number}행 {field_key}: 올바른 날짜를 입력해 주세요.")
        if errors:
            raise ValueError("\n".join(errors))

    def _safe_filename(self, title):
        invalid_chars = '<>:"/\\|?*'
        if title.lower().endswith(".pdf"):
            title = title[:-4]
        safe = ''.join("_" if ch in invalid_chars else ch for ch in title).strip().rstrip(".")
        return safe or self.file_path.stem

    def _resolve_output_pdf_path(self, folder, title, ask_on_duplicate):
        safe_title = self._safe_filename(title)
        candidate = folder / f"{safe_title}.pdf"
        if not candidate.exists():
            return candidate

        numbered = self._next_numbered_path(folder, safe_title)
        if ask_on_duplicate:
            answer = QMessageBox.question(
                self,
                "중복 파일명",
                f"같은 제목의 파일이 이미 있습니다.\n{numbered.name} 이름으로 저장할까요?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if answer != QMessageBox.Yes:
                return None
        return numbered

    def _next_numbered_path(self, folder, safe_title):
        number = 1
        while True:
            candidate = folder / f"{safe_title}({number}).pdf"
            if not candidate.exists():
                return candidate
            number += 1

    def _ensure_docx2pdf(self):
        docx2pdf = self._load_docx2pdf()
        if docx2pdf is not None:
            return docx2pdf

        install_answer = QMessageBox.question(
            self,
            'docx2pdf 설치',
            'docx2pdf 패키지가 설치되어 있지 않습니다. 설치하시겠습니까?',
            QMessageBox.Yes | QMessageBox.No,
        )
        if install_answer != QMessageBox.Yes:
            log("DocumentTool", "docx2pdf 설치 거부됨; PDF 생성 취소", level="WARNING")
            raise RuntimeError("docx2pdf 라이브러리가 설치되지 않았습니다.")

        return self._install_docx2pdf()

    def _apply_values_to_doc(self, doc, values_dict):
        values_dict = dict(values_dict)

        def resolve_date(field_key, visited=None):
            visited = set(visited or ())
            if field_key in visited:
                return None
            visited.add(field_key)
            setting = self.settings.get("fields", {}).get(field_key, {})
            if setting.get("type") != "date":
                return None
            if setting.get("date_default_type", "today") != "field_calculation":
                return self._parse_date_or_none(values_dict.get(field_key, "")) or QDate.currentDate()
            source_date = resolve_date(setting.get("source_date_field", ""), visited)
            if not source_date:
                return None
            sign = int(setting.get("date_offset_sign", 1) or 1)
            days = int(setting.get("date_offset_days", 0) or 0)
            return source_date.addDays(sign * days)

        for field_key, setting in self.settings.get("fields", {}).items():
            if setting.get("type") == "date" and setting.get("date_default_type") == "field_calculation":
                calculated_date = resolve_date(field_key)
                values_dict[field_key] = (
                    calculated_date.toString(setting.get("date_format", "yyyy-MM-dd"))
                    if calculated_date
                    else ""
                )
        for field_key, setting in self.settings.get("fields", {}).items():
            if setting.get("type") == "folder":
                values_dict[field_key] = self._folder_image_value(
                    setting,
                    values_dict.get(setting.get("value_field", ""), ""),
                )
        sdt_list = doc.element.findall('.//' + qn('w:sdt'))
        for idx, sdt in enumerate(sdt_list):
            info = self._content_control_info(sdt, idx)
            if not info:
                continue
            field_key = info["tag"] or info["alias"] or f"필드 {idx + 1}"
            if field_key not in values_dict:
                continue

            new_value = values_dict[field_key]
            sdt_content = sdt.find(qn('w:sdtContent'))
            if sdt_content is None:
                continue

            if isinstance(new_value, dict) and new_value.get("type") == "image":
                self._set_sdt_image(doc, sdt_content, new_value)
                continue

            t_list = sdt_content.findall('.//' + qn('w:t'))
            if t_list:
                t_list[0].text = str(new_value)
                for t in t_list[1:]:
                    parent = t.getparent()
                    try:
                        parent.remove(t)
                    except Exception:
                        pass
            else:
                run = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = str(new_value)
                run.append(t)
                sdt_content.append(run)

    @staticmethod
    def _folder_image_value(setting, source_value):
        folder = Path(setting.get("folder_path", ""))
        target_name = str(source_value or "").strip()
        if not target_name or not folder.is_dir():
            return ""
        target_name = target_name.casefold()
        matches = sorted(
            (
                path for path in folder.iterdir()
                if path.is_file()
                and path.suffix.casefold() in IMAGE_EXTENSIONS
                and path.stem.casefold() == target_name
            ),
            key=lambda path: path.name.casefold(),
        )
        if not matches:
            return ""
        return {
            "type": "image",
            "path": str(matches[0]),
            "width": int(setting.get("width", 30) or 30),
            "height": int(setting.get("height", 30) or 30),
        }

    def _set_sdt_image(self, doc, sdt_content, value):
        image_path = Path(value.get("path", ""))
        if not value.get("path") or not image_path.exists():
            return

        for child in list(sdt_content):
            sdt_content.remove(child)

        # 문단 안에 위치한 run-level 콘텐츠 컨트롤에는 w:r만 들어갈 수 있다.
        # 여기에 w:p를 넣으면 문단 안에 문단이 중첩되어 Word가 DOCX를 손상된
        # 파일로 판단한다. block-level 컨트롤일 때만 새 문단을 만든다.
        parent = sdt_content.getparent().getparent()
        if parent is not None and parent.tag == qn('w:p'):
            run_element = OxmlElement('w:r')
            sdt_content.append(run_element)
            run = Run(run_element, doc)
        else:
            paragraph_element = OxmlElement('w:p')
            sdt_content.append(paragraph_element)
            paragraph = Paragraph(paragraph_element, doc)
            run = paragraph.add_run()
        run.add_picture(
            str(image_path),
            width=Mm(int(value.get("width", 30) or 30)),
            height=Mm(int(value.get("height", 30) or 30)),
        )
