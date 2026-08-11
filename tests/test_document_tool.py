from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ui.pages.document_tool import ContentControlDialog, branch_control_parts


def test_branch_control_name_parsing():
    assert branch_control_parts("분기_시설종류_요양원") == ("시설종류", "요양원")
    assert branch_control_parts("분기_시설종류_주야간_보호") == ("시설종류", "주야간_보호")
    assert branch_control_parts("시설종류_요양원") is None
    assert branch_control_parts("분기__요양원") is None


def test_branch_generation_only_updates_selected_control():
    dialog = ContentControlDialog.__new__(ContentControlDialog)
    combo = lambda value: SimpleNamespace(currentData=lambda: value)
    dialog.generate_widgets = {
        "분기_변경전_대표자": {
            "type_combo": combo("branch"),
            "branch_option_combo": combo("분기_변경전_시설명칭"),
            "branch_members": [
                "분기_변경전_대표자",
                "분기_변경전_시설명칭",
                "분기_변경전_시설소재지",
            ],
        }
    }
    dialog._field_value = lambda field_key: f"입력:{field_key}"

    assert dialog._generation_values() == {
        "분기_변경전_시설명칭": "입력:분기_변경전_시설명칭"
    }


def test_inline_image_control_contains_run_instead_of_nested_paragraph():
    paragraph = OxmlElement("w:p")
    content_control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    content_control.append(content)
    paragraph.append(content_control)

    value = {
        "type": "image",
        "path": str(Path(__file__).resolve()),
        "width": 15,
        "height": 15,
    }

    with patch("ui.pages.document_tool.Run.add_picture") as add_picture:
        ContentControlDialog._set_sdt_image(object(), object(), content, value)

    assert [child.tag for child in content] == [qn("w:r")]
    add_picture.assert_called_once()


def test_folder_image_value_matches_filename_without_extension(tmp_path):
    matching_image = tmp_path / "홍길동.PNG"
    matching_image.write_bytes(b"image")
    (tmp_path / "홍길동.txt").write_text("not an image", encoding="utf-8")

    value = ContentControlDialog._folder_image_value(
        {"folder_path": str(tmp_path)},
        "홍길동",
    )

    assert value == {
        "type": "image",
        "path": str(matching_image),
        "width": 30,
        "height": 30,
    }


def test_folder_image_value_returns_empty_when_image_is_missing(tmp_path):
    value = ContentControlDialog._folder_image_value(
        {"folder_path": str(tmp_path)},
        "없는 이름",
    )

    assert value == ""


def test_folder_image_value_uses_configured_size(tmp_path):
    matching_image = tmp_path / "김영희.jpg"
    matching_image.write_bytes(b"image")

    value = ContentControlDialog._folder_image_value(
        {"folder_path": str(tmp_path), "width": 45, "height": 20},
        "김영희",
    )

    assert value["width"] == 45
    assert value["height"] == 20


def test_prepare_excel_row_converts_image_field_path(tmp_path):
    image_path = tmp_path / "stamp.png"
    image_path.write_bytes(b"image")
    dialog = ContentControlDialog.__new__(ContentControlDialog)
    dialog.settings = {
        "fields": {
            "image_test": {
                "type": "branch_value",
                "branch_value_key": "custom_법인인감",
            }
        }
    }
    dialog.branch_fields = [
        {"key": "custom_법인인감", "label": "법인인감", "type": "image"}
    ]
    dialog.root = tmp_path
    dialog.field_widgets = {}

    values = dialog._prepare_excel_row_values(
        {"image_test": str(image_path)},
        row_number=2,
    )

    assert values["image_test"] == {
        "type": "image",
        "path": str(image_path),
        "width": 30,
        "height": 30,
    }


def test_prepare_excel_row_keeps_text_field_as_text(tmp_path):
    dialog = ContentControlDialog.__new__(ContentControlDialog)
    dialog.settings = {"fields": {"memo": {"type": "text"}}}
    dialog.branch_fields = []
    dialog.root = tmp_path
    dialog.field_widgets = {}

    values = dialog._prepare_excel_row_values({"memo": "C:\\image.png"}, row_number=2)

    assert values["memo"] == "C:\\image.png"


def test_prepare_excel_row_rejects_missing_image(tmp_path):
    dialog = ContentControlDialog.__new__(ContentControlDialog)
    dialog.settings = {
        "fields": {
            "seal": {
                "type": "branch_value",
                "branch_value_key": "custom_법인인감",
            }
        }
    }
    dialog.branch_fields = [{"key": "custom_법인인감", "type": "image"}]
    dialog.root = tmp_path
    dialog.field_widgets = {}

    try:
        dialog._prepare_excel_row_values({"seal": "missing.png"}, row_number=3)
    except ValueError as error:
        assert "3행 seal" in str(error)
        assert "이미지 파일을 찾을 수 없습니다" in str(error)
    else:
        raise AssertionError("존재하지 않는 이미지 경로는 오류여야 합니다.")
